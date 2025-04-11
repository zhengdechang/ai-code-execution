import asyncio
import importlib
from kernel_client import KernelClient


async def main():
    async with KernelClient("http://localhost:8888") as client:
        code4 = """
        from docx import Document
        import numpy as np
        
        # 创建一个 Word 文档
        doc = Document()
        doc.add_heading('随机数据报告', level=1)
        
        # 创建示例数据
        data = {
            'A': np.random.normal(0, 1, 10),
            'B': np.random.normal(5, 2, 10),
            'C': np.random.uniform(1, 10, 10)
        }
        
        # 添加表格
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '索引'
        hdr_cells[1].text = 'A'
        hdr_cells[2].text = 'B'
        hdr_cells[3].text = 'C'
        
        for i in range(10):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = f"{data['A'][i]:.2f}"
            row_cells[2].text = f"{data['B'][i]:.2f}"
            row_cells[3].text = f"{data['C'][i]:.2f}"
        
        # 保存文档
        doc.save('document.docx')
        print("文档已保存为 document.docx")
        """

        # 检查缺失的包
        required_packages = ['python-docx', 'numpy']
        install_result = await asyncio.wait_for(client.install_packages(required_packages), timeout=10)
        if install_result.success:
            print("\ninstall packages:", install_result.output)
        else:
            print("错误:", install_result.error)

        # 执行代码
        result4 = await asyncio.wait_for(client.execute_code(code4), timeout=10)
        print("\n示例 4 结果:")
        if result4.success:
            print(result4.output)
        else:
            print("错误:", result4.error)

if __name__ == "__main__":
    asyncio.run(main())
